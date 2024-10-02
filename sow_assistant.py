import streamlit as st
from google.oauth2 import service_account
from google.cloud import storage
from langchain_google_vertexai import VertexAI
from langchain_google_vertexai.embeddings import VertexAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.schema import Document
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from docx import Document as DocxDocument

# Create credentials object
credentials = service_account.Credentials.from_service_account_info(
    st.secrets["gcp_service_account"]
)

def load_documents_from_gcs(bucket_name):
    client = storage.Client(credentials=credentials)
    bucket = client.get_bucket(bucket_name)
    documents = []
    blobs = bucket.list_blobs()

    for blob in blobs:
        if blob.name.endswith('.txt'):
            content = blob.download_as_text()
            metadata = {"source": blob.name}
            documents.append(Document(page_content=content, metadata=metadata))

    return documents

@st.cache_resource
def initialize_vectorstore(bucket_name):
    documents = load_documents_from_gcs(bucket_name)
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    splits = text_splitter.split_documents(documents)
    embedding_model = VertexAIEmbeddings("textembedding-gecko", credentials=credentials)
    vectorstore = FAISS.from_documents(documents=splits, embedding=embedding_model)
    return vectorstore

class SOWAssistant:
    def __init__(self, bucket_name):
        self.llm = VertexAI(model_name="gemini-1.0-pro-002", temperature=0.3, credentials=credentials)
        self.embedding_model = VertexAIEmbeddings("textembedding-gecko", credentials=credentials)
        self.vectorstore = initialize_vectorstore(bucket_name)
        self.sow_content = ""

        self.sow_template = """
        You are an expert in creating Statements of Work (SOWs) for various projects.
        Using the information provided and the retrieved similar SOWs,
        Please create a detailed SOW that includes only these following section:
        1. Project name
        2. Client name
        3. Delivered by: IM digital
        4. Project overview and objective
        5. Project goals
        6. Scope of work
        7. Deliverables
        8. Timeline
        9. Cost estimation (investment)
        10. Payment terms

        Consider the following information:
        {question}

        Use the following similar SOWs as reference, but tailor the new SOW to the specific requirements provided:
        {context}

        Ensure that the generated SOW is tailored to the specific project details provided in the user input,
        while leveraging relevant information from the similar SOWs.
        Note: Strictly follow the format, don't add any extra information. If you don't have any similar SOWs, please provide an empty string.
        """

        self.customize_template = """
        You are an expert in customizing Statements of Work (SOWs) for various projects.
        Using the provided information and the retrieved similar SOWs, update the SOW as requested.

        Current SOW:
        {current_sow}

        Customization Request:
        {customization_request}

        Consider the following similar SOWs as reference for your customization:
        {context}

        Generate an updated SOW that incorporates the requested changes while maintaining the overall structure and quality of the document. Explain the changes you've made.
        """

        self.sow_prompt = PromptTemplate(template=self.sow_template, input_variables=["question", "context"])
        self.customize_prompt = PromptTemplate(template=self.customize_template, input_variables=["current_sow", "customization_request", "context"])

        self.qa_chain = RetrievalQA.from_chain_type(
            llm=self.llm,
            chain_type="stuff",
            retriever=self.vectorstore.as_retriever(search_kwargs={"k": 10}),
            return_source_documents=True,
            chain_type_kwargs={"prompt": self.sow_prompt}
        )

    def generate_sow(self, query):
        similar_docs = self.vectorstore.similarity_search(query, k=3)
        context = "\n\n".join([doc.page_content for doc in similar_docs])

        result = self.qa_chain.invoke({"query": query})
        self.sow_content = result["result"]

        sources = [
            {
                "content": doc.page_content,
                "source": doc.metadata.get("source", "Unknown Source"),
                "similarity": getattr(doc, 'similarity', None)
            }
            for doc in result["source_documents"]
        ]

        return {
            "sow": self.sow_content,
            "sources": sources,
            "context": context
        }

    def customize_sow(self, customization_request):
        combined_query = f"""
        Current SOW:
        {self.sow_content}
        Customization Request:
        {customization_request}
        """
        customize_chain = RetrievalQA.from_chain_type(
            llm=self.llm,
            chain_type="stuff",
            retriever=self.vectorstore.as_retriever(search_kwargs={"k": 10}),
            return_source_documents=True,
            chain_type_kwargs={
                "prompt": PromptTemplate(
                    template=self.customize_template,
                    input_variables=["question", "context"]
                )
            }
        )

        result = customize_chain.invoke({"query": combined_query})
        self.sow_content = result["result"]

        sources = [
            {
                "content": doc.page_content,
                "source": doc.metadata.get("source", "Unknown Source"),
                "similarity": getattr(doc, 'similarity', None)
            }
            for doc in result["source_documents"]
        ]
        return {
            "sow": self.sow_content,
            "sources": sources
        }

    def export_sow_docx(self, export_prompt):
        formatting_prompt = f"""
        Format the following SOW content for export to a DOCX file.
        Organize the content into clear sections with appropriate headings.
        Use bullet points or numbered lists where applicable.
        Ensure the formatting is clean and professional.

        SOW Content:
        {self.sow_content}

        User's export request:
        {export_prompt}

        Provide the formatted content as plain text, using markdown-style formatting
        for headings (e.g., # for main headings, ## for subheadings) and standard
        markdown for lists and emphasis.
        """

        formatted_content = self.llm.predict(formatting_prompt)

        doc = DocxDocument()

        for line in formatted_content.split('\n'):
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('1. '):
                doc.add_paragraph(line[3:], style='List Number')
            else:
                doc.add_paragraph(line)

        filename = "generated_sow.docx"
        doc.save(filename)

        return f"SOW has been exported to {filename}"
